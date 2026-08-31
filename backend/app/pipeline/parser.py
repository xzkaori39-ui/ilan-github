"""文档格式解析层：PDF / Word / Markdown / HTML / TXT 统一解析为结构化块。

保留标题层级、表格、列表结构。可选依赖按需延迟导入，缺失时降级。
"""
from __future__ import annotations

import re
import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional, Union

from app.utils.logging import get_logger

logger = get_logger(__name__)

# 仅匹配“章/节/部分/篇/编”等标题，不匹配“条/款”（条款属于正文，交由切片器处理）
HEADING_RE = re.compile(r"^(第[一二三四五六七八九十百0-9]+[章节部分篇编]|[一二三四五六七八九十]+、|[0-9]+[.、．]\s*[^0-9])")
LARGE_PDF_THRESHOLD_BYTES = 32 * 1024 * 1024


@dataclass
class Block:
    type: str  # heading | paragraph | list_item | table
    level: int  # 标题层级 1-6
    text: str
    page: Optional[int] = None
    extra: dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedDocument:
    title: str
    blocks: list[Block]
    raw_text: str = ""
    page_count: int = 0
    meta: dict[str, Any] = field(default_factory=dict)

    @property
    def text(self) -> str:
        if self.raw_text:
            return self.raw_text
        return "\n".join(b.text for b in self.blocks if b.text)


class DocumentParser:
    """按文件类型分派解析器。"""

    def parse(self, path: Union[str, Path]) -> ParsedDocument:
        path = Path(path)
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._parse_pdf(path)
        if suffix == ".docx":
            return self._parse_docx(path)
        if suffix in (".md", ".markdown"):
            return self._parse_markdown(path)
        if suffix in (".html", ".htm"):
            return self._parse_html(path)
        if suffix in (".txt", ""):
            return self._parse_text(path)
        raise ValueError(f"不支持的文件类型: {suffix}")

    # ---------- PDF ----------
    def _parse_pdf(self, path: Path) -> ParsedDocument:
        text_pages: list[str] = []
        ocr_used = False
        if path.stat().st_size >= LARGE_PDF_THRESHOLD_BYTES:
            logger.info("大体积 PDF 使用 pypdf 低内存解析: %s", path.name)
            page_count = self._pdf_page_count(path)
            if self._pdf_has_text_layer(path, page_count):
                text_pages = self._parse_pdf_pypdf(path)
            else:
                logger.info("大体积 PDF 无文本层，直接逐页 OCR: %s", path.name)
                text_pages = self._ocr_pdf(path, page_count)
                ocr_used = True
        else:
            try:
                import pdfplumber  # 延迟导入

                with pdfplumber.open(str(path)) as pdf:
                    for page in pdf.pages:
                        text_pages.append(page.extract_text() or "")
            except ImportError:
                text_pages = self._parse_pdf_pypdf(path)
            except Exception as exc:  # noqa: BLE001
                logger.warning("pdfplumber 解析失败(%s)，回退 pypdf", exc)
                text_pages = self._parse_pdf_pypdf(path)

        if not ocr_used and text_pages and not any(page.strip() for page in text_pages):
            logger.info("PDF 无文本层，使用本地 OCR 解析扫描件: %s", path.name)
            text_pages = self._ocr_pdf(path, len(text_pages))
            ocr_used = True

        blocks = self._blocks_from_pages(text_pages)
        title = path.stem
        return ParsedDocument(
            title=title,
            blocks=blocks,
            raw_text="\n".join(text_pages),
            page_count=len(text_pages),
            meta={"ocr_used": ocr_used},
        )

    def _parse_pdf_pypdf(self, path: Path) -> list[str]:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        return [(page.extract_text() or "") for page in reader.pages]

    def _pdf_page_count(self, path: Path) -> int:
        """读取 PDF 页数，不将整份文档展开为 pypdf 页面对象。"""
        try:
            result = subprocess.run(
                ["pdfinfo", str(path)], check=True, capture_output=True, text=True, timeout=30,
            )
            match = re.search(r"^Pages:\s+(\d+)\s*$", result.stdout, re.MULTILINE)
            if match:
                return int(match.group(1))
        except (FileNotFoundError, subprocess.SubprocessError) as exc:
            logger.warning("pdfinfo 读取页数失败(%s)，回退 pypdf: %s", exc, path.name)

        from pypdf import PdfReader

        return len(PdfReader(str(path)).pages)

    def _pdf_has_text_layer(self, path: Path, page_count: int) -> bool:
        """抽样页检测文本层；扫描件无需先经 pypdf 全量展开。"""
        sample_pages = sorted({1, max(1, page_count // 2), page_count})
        for page_no in sample_pages:
            try:
                result = subprocess.run(
                    ["pdftotext", "-f", str(page_no), "-l", str(page_no), str(path), "-"],
                    check=True,
                    capture_output=True,
                    text=True,
                    timeout=30,
                )
                if result.stdout.strip():
                    return True
            except FileNotFoundError as exc:
                logger.warning("pdftotext 不可用，回退 pypdf: %s", path.name)
                return True
            except (subprocess.SubprocessError, UnicodeDecodeError) as exc:
                logger.warning("PDF 文本层探测失败(%s)，回退 pypdf: %s", exc, path.name)
                return True
        return False

    def _ocr_pdf(self, path: Path, page_count: int) -> list[str]:
        """逐页渲染并识别扫描 PDF，避免一次性把整本手册加载进内存。"""
        try:
            from rapidocr_onnxruntime import RapidOCR
        except ImportError as exc:
            raise RuntimeError("扫描 PDF 需要 rapidocr-onnxruntime，请安装后重试") from exc

        ocr = RapidOCR()
        text_pages: list[str] = []
        with tempfile.TemporaryDirectory(prefix="ilan-ocr-") as tmp_dir:
            tmp_path = Path(tmp_dir)
            for page_no in range(1, page_count + 1):
                image_prefix = tmp_path / f"page-{page_no}"
                image_path = image_prefix.with_suffix(".png")
                try:
                    subprocess.run(
                        [
                            "pdftoppm", "-f", str(page_no), "-l", str(page_no),
                            "-r", "150", "-png", "-singlefile", str(path), str(image_prefix),
                        ],
                        check=True,
                        capture_output=True,
                        text=True,
                        timeout=120,
                    )
                except FileNotFoundError as exc:
                    raise RuntimeError("扫描 PDF OCR 依赖 pdftoppm，请安装 poppler-utils") from exc
                except subprocess.CalledProcessError as exc:
                    raise RuntimeError(f"第 {page_no} 页 PDF 渲染失败: {exc.stderr.strip()}") from exc
                except subprocess.TimeoutExpired as exc:
                    raise RuntimeError(f"第 {page_no} 页 PDF 渲染超时") from exc

                result, _ = ocr(str(image_path))
                text_pages.append("\n".join(line[1] for line in (result or [])))
        return text_pages

    # ---------- Word ----------
    def _parse_docx(self, path: Path) -> ParsedDocument:
        import docx  # 延迟导入

        doc = docx.Document(str(path))
        blocks: list[Block] = []
        # 文档内段落遍历（含表格）
        from docx.document import Document as _Doc
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        title = path.stem
        body = doc.element.body
        for child in body.iterchildren():
            if child.tag.endswith("}p"):
                para = Paragraph(child, doc)
                blocks.append(self._paragraph_to_block(para))
            elif child.tag.endswith("}tbl"):
                table = Table(child, doc)
                blocks.append(self._table_to_block(table))
        raw = "\n".join(b.text for b in blocks if b.text)
        return ParsedDocument(title=title, blocks=blocks, raw_text=raw, meta={"format": "docx"})

    def _paragraph_to_block(self, para: Any) -> Block:
        text = para.text.strip()
        style_name = (para.style.name if para.style and para.style.name else "") or ""
        level = 0
        if "Heading" in style_name or "标题" in style_name:
            digits = re.findall(r"\d+", style_name)
            level = int(digits[0]) if digits else 1
            return Block(type="heading", level=level, text=text)
        if text.startswith(("•", "- ", "·", "1.", "2.", "3.", "（", "(")):
            return Block(type="list_item", level=1, text=text)
        return Block(type="paragraph", level=0, text=text)

    def _table_to_block(self, table: Any) -> Block:
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            rows.append(" | ".join(cells))
        return Block(type="table", level=0, text="\n".join(rows))

    # ---------- Markdown ----------
    def _parse_markdown(self, path: Path) -> ParsedDocument:
        text = path.read_text(encoding="utf-8", errors="ignore")
        lines = text.splitlines()
        blocks: list[Block] = []
        title = path.stem
        buf: list[str] = []
        buf_type = "paragraph"

        def flush() -> None:
            nonlocal buf
            if buf:
                blocks.append(Block(type=buf_type, level=0, text="\n".join(buf)))
                buf = []

        for line in lines:
            m = re.match(r"^(#{1,6})\s+(.*)$", line)
            if m:
                flush()
                blocks.append(Block(type="heading", level=len(m.group(1)), text=m.group(2).strip()))
                continue
            if re.match(r"^\s*[-*+]\s+", line):
                flush()
                buf_type = "list_item"
                buf.append(re.sub(r"^\s*[-*+]\s+", "", line))
                continue
            if line.strip() == "":
                flush()
                buf_type = "paragraph"
                continue
            buf.append(line)
        flush()
        return ParsedDocument(title=title, blocks=blocks, raw_text=text, meta={"format": "markdown"})

    # ---------- HTML ----------
    def _parse_html(self, path: Path) -> ParsedDocument:
        from bs4 import BeautifulSoup  # 延迟导入

        html = path.read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(html, "html.parser")
        title = soup.title.string if soup.title and soup.title.string else path.stem
        blocks: list[Block] = []
        for tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "table"]):
            name = tag.name
            text = tag.get_text(" ", strip=True)
            if not text:
                continue
            if name.startswith("h"):
                blocks.append(Block(type="heading", level=int(name[1]), text=text))
            elif name == "li":
                blocks.append(Block(type="list_item", level=1, text=text))
            elif name == "table":
                blocks.append(Block(type="table", level=0, text=text))
            else:
                blocks.append(Block(type="paragraph", level=0, text=text))
        return ParsedDocument(title=title, blocks=blocks, raw_text=soup.get_text(" ", strip=True), meta={"format": "html"})

    # ---------- TXT ----------
    def _parse_text(self, path: Path) -> ParsedDocument:
        text = path.read_text(encoding="utf-8", errors="ignore")
        blocks = self._blocks_from_pages([text])
        return ParsedDocument(title=path.stem, blocks=blocks, raw_text=text, page_count=1)

    # ---------- 辅助 ----------
    def _blocks_from_pages(self, pages: list[str]) -> list[Block]:
        blocks: list[Block] = []
        for page_no, text in enumerate(pages, start=1):
            for line in text.splitlines():
                line = line.strip()
                if not line:
                    continue
                m = HEADING_RE.match(line)
                if m and len(line) <= 60:
                    blocks.append(Block(type="heading", level=2, text=line, page=page_no))
                else:
                    blocks.append(Block(type="paragraph", level=0, text=line, page=page_no))
        return blocks
